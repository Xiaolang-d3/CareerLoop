from __future__ import annotations

import unittest
from io import BytesIO

from docx import Document

from app.resume.parser import normalize_resume_text, parse_resume
from app.profile.service import parse_candidate_resume


class ResumeParserTest(unittest.TestCase):
    def test_parses_utf8_text_resume(self) -> None:
        text = parse_resume("resume.txt", "张三\nPython 开发工程师\n五年后端开发经验".encode())
        self.assertIn("Python 开发工程师", text)

    def test_parse_response_includes_profile_fill_suggestions(self) -> None:
        result = parse_candidate_resume(
            "resume.txt",
            "姓名：张三\n求职意向：后端工程师\n期望城市：上海\n技能：Python、Docker".encode(),
            "fast",
        )

        self.assertEqual(result["suggested_profile"]["name"], "张三")
        self.assertEqual(result["suggested_profile"]["target_roles"], ["后端工程师"])
        self.assertEqual(result["suggested_profile"]["target_cities"], ["上海"])
        self.assertGreaterEqual(set(result["suggested_profile"]["skills"]), {"Python", "Docker"})

    def test_parses_docx_paragraphs_and_tables(self) -> None:
        document = Document()
        document.add_heading("个人简历", level=1)
        document.add_paragraph("负责 FastAPI 和 Agent 平台开发")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "技能"
        table.cell(0, 1).text = "Python"
        stream = BytesIO()
        document.save(stream)

        text = parse_resume("resume.docx", stream.getvalue())

        self.assertIn("Agent 平台开发", text)
        self.assertIn("技能\tPython", text)

    def test_normalizes_private_use_bullets_and_invisible_characters(self) -> None:
        text = parse_resume(
            "resume.txt",
            "智能会议总结\n\uf0b7 基于 LangChain 搭建统一网关\u200b\n".encode(),
        )

        self.assertIn("- 基于 LangChain 搭建统一网关", text)
        self.assertNotIn("\uf0b7", text)
        self.assertNotIn("\u200b", text)

    def test_unwraps_pdf_visual_line_breaks_without_merging_headings(self) -> None:
        text = normalize_resume_text(
            "个人优势\n"
            "- 熟悉 Python / FastAPI / Docker 等技术栈，\n"
            "擅长实时语音链路与多模型协同。\n"
            "- 独立完成业务内容生产链路，有效提升业务内容产出效\n"
            "率 60%+\n"
            "工作经历\n"
            "某公司\n"
            "AI 应用工程师\n"
        )

        self.assertIn("Docker 等技术栈，擅长实时语音链路与多模型协同。", text)
        self.assertIn("产出效率 60%+", text)
        self.assertIn("\n工作经历\n", text)
        self.assertIn("\n某公司\n", text)
        self.assertIn("AI 应用工程师", text)
        self.assertNotIn("某公司AI 应用工程师", text)

    def test_splits_jammed_contact_and_certificate_fields(self) -> None:
        text = normalize_resume_text(
            "小程\n"
            "邮箱: [邮箱已隐藏] 英语: CET-6GitHub: https://github.com/Xiaolang-d3\n"
            "求职意向：后端工程师\n"
        )

        self.assertIn("邮箱: [邮箱已隐藏]", text)
        self.assertIn("英语: CET-6", text)
        self.assertIn("GitHub: https://github.com/Xiaolang-d3", text)
        self.assertNotIn("CET-6GitHub", text)
        self.assertIn("\n求职意向：后端工程师", text)

    def test_rejects_unsupported_or_empty_files(self) -> None:
        with self.assertRaises(ValueError):
            parse_resume("resume.pages", b"not supported")
        with self.assertRaises(ValueError):
            parse_resume("resume.txt", b"")


if __name__ == "__main__":
    unittest.main()
