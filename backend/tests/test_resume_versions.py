from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document
from pypdf import PdfReader

from app.db import connect, init_db, json_dump
from app.jobs.service import create_job, update_job
from evaluation_helpers import seed_confirmed_facts_and_evaluation
from app.resume.layout import (
    compose_rendered_sections,
    contact_link_target,
    split_document_name,
    split_entry_heading,
    split_resume_layout,
)
from app.resume.versions import (
    create_baseline_resume_version,
    create_resume_version,
    export_resume_version,
    get_resume_version,
    list_resume_versions,
    update_resume_change,
    update_resume_version,
)


class ResumeVersionTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = TemporaryDirectory()
        self.db_path = Path(self.temp_dir.name) / "resume-versions.db"
        init_db(self.db_path)
        with connect(self.db_path) as conn:
            cursor = conn.execute(
                """
                INSERT INTO profiles (
                    name, resume_text, resume_redacted_text, privacy_mode, skills_json
                ) VALUES (?, ?, ?, 'redacted', ?)
                """,
                (
                    "测试用户",
                    "13800138000\n5年产品经验\n负责 Agent 产品规划和需求分析。\n使用 Python 与 FastAPI 完成内部原型。",
                    "[手机号已隐藏]\n5年产品经验\n负责 Agent 产品规划和需求分析。\n使用 Python 与 FastAPI 完成内部原型。",
                    json_dump(["Agent", "Python", "FastAPI"]),
                ),
            )
            conn.execute(
                "INSERT INTO preferences (profile_id) VALUES (?)",
                (cursor.lastrowid,),
            )
        self.job = create_job(
            {
                "job_title": "AI 产品经理",
                "company_name": "示例科技",
                "description": (
                    "负责 Agent 产品规划和需求分析；"
                    "要求熟悉 Python 与 FastAPI；"
                    "能够推动产品原型落地。"
                ),
            },
            self.db_path,
        )
        seed_confirmed_facts_and_evaluation(self.job["id"], self.db_path)

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_version_is_persistent_traceable_and_does_not_overwrite_profile(self) -> None:
        version = create_resume_version(self.job["id"], self.db_path)

        self.assertEqual(version["status"], "draft")
        self.assertGreaterEqual(version["change_count"], 3)
        self.assertEqual(version["change_counts"]["pending"], version["change_count"])
        self.assertIn("AI 产品经理", version["rendered_content"])
        self.assertIn("Python", version["rendered_content"])
        self.assertNotIn("13800138000", json_dump(version))
        self.assertTrue(
            all(change["evidence"] for change in version["changes"])
        )
        self.assertEqual(
            list_resume_versions(self.job["id"], self.db_path)[0]["id"],
            version["id"],
        )
        self.assertNotIn(
            "base_content",
            list_resume_versions(self.job["id"], self.db_path)[0],
        )
        with connect(self.db_path) as conn:
            profile = conn.execute("SELECT resume_text FROM profiles").fetchone()
        self.assertIn("13800138000", profile["resume_text"])

    def test_change_decisions_and_user_edits_rebuild_preview(self) -> None:
        version = create_resume_version(self.job["id"], self.db_path)
        summary = next(
            change for change in version["changes"] if change["section_key"] == "summary"
        )

        rejected = update_resume_change(
            version["id"],
            summary["id"],
            decision="rejected",
            db_path=self.db_path,
        )
        self.assertNotIn("简历中可验证的岗位相关能力", rejected["rendered_content"])
        self.assertEqual(rejected["change_counts"]["rejected"], 1)

        edited = update_resume_change(
            version["id"],
            summary["id"],
            decision="accepted",
            after_text="## 职业概述\n这是我本人确认并编辑的职业概述。",
            db_path=self.db_path,
        )
        edited_change = next(
            change for change in edited["changes"] if change["id"] == summary["id"]
        )
        self.assertEqual(edited_change["user_edited"], 1)
        self.assertIn("本人确认", edited["rendered_content"])

        final = update_resume_version(
            version["id"],
            status="final",
            db_path=self.db_path,
        )
        self.assertEqual(final["status"], "final")
        self.assertEqual(
            get_resume_version(version["id"], self.db_path)["status"],
            "final",
        )
        reopened = update_resume_change(
            version["id"],
            summary["id"],
            decision="pending",
            db_path=self.db_path,
        )
        self.assertEqual(reopened["status"], "draft")

    def test_template_is_persisted_per_resume_version(self) -> None:
        version = create_resume_version(self.job["id"], self.db_path)

        self.assertEqual(version["template_id"], "classic")
        updated = update_resume_version(
            version["id"],
            template_id="compact",
            db_path=self.db_path,
        )

        self.assertEqual(updated["template_id"], "compact")
        self.assertEqual(
            get_resume_version(version["id"], self.db_path)["template_id"],
            "compact",
        )
        self.assertEqual(
            list_resume_versions(self.job["id"], self.db_path)[0]["template_id"],
            "compact",
        )
        self.assertEqual(version["style_id"], "navy")
        styled = update_resume_version(
            version["id"],
            style_id="wine",
            db_path=self.db_path,
        )
        self.assertEqual(styled["style_id"], "wine")
        self.assertEqual(
            get_resume_version(version["id"], self.db_path)["style_id"],
            "wine",
        )
        self.assertEqual(version["layout"], {"spacing": 100, "one_page": False})
        packed = update_resume_version(
            version["id"],
            layout={"spacing": 80, "one_page": True},
            db_path=self.db_path,
        )
        self.assertEqual(packed["layout"], {"spacing": 80, "one_page": True})

    def test_docx_and_pdf_exports_are_valid_documents(self) -> None:
        version = create_resume_version(self.job["id"], self.db_path)

        docx_bytes, docx_name, docx_type = export_resume_version(
            version["id"],
            "docx",
            self.db_path,
        )
        self.assertTrue(docx_name.endswith(".docx"))
        self.assertIn("wordprocessingml", docx_type)
        document = Document(BytesIO(docx_bytes))
        docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("AI 产品经理", docx_text)
        self.assertNotIn("13800138000", docx_text)

        pdf_bytes, pdf_name, pdf_type = export_resume_version(
            version["id"],
            "pdf",
            self.db_path,
        )
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))
        self.assertTrue(pdf_name.endswith(".pdf"))
        self.assertEqual(pdf_type, "application/pdf")
        self.assertGreaterEqual(len(PdfReader(BytesIO(pdf_bytes)).pages), 1)

    def test_pdf_export_renders_chinese_with_a_cjk_font(self) -> None:
        # reportlab 的内置 Latin 字体会把中文渲染成方块，导出必须挂上 CJK 字体。
        version = create_resume_version(self.job["id"], self.db_path)

        for style_id in ("navy", "ink"):
            update_resume_version(version["id"], style_id=style_id, db_path=self.db_path)
            pdf_bytes, _name, _type = export_resume_version(
                version["id"],
                "pdf",
                self.db_path,
            )
            page = PdfReader(BytesIO(pdf_bytes)).pages[0]
            with self.subTest(style_id=style_id):
                self.assertIn("产品经理", page.extract_text())
                base_fonts = {
                    str(font.get("/BaseFont", ""))
                    for font in (page.get("/Resources", {}).get("/Font", {}) or {}).values()
                }
                self.assertTrue(
                    any(
                        "STSong" in base_font or "+" in base_font
                        for base_font in base_fonts
                    ),
                    f"{style_id} 没有使用 CID 或内嵌的中文字体：{sorted(base_fonts)}",
                )

    def test_material_uses_confirmed_facts_instead_of_legacy_resume_text(self) -> None:
        with connect(self.db_path) as conn:
            conn.execute("UPDATE profiles SET resume_redacted_text = ''")
        version = create_resume_version(self.job["id"], self.db_path)
        self.assertIn("Agent 产品规划", version["base_content"])
        self.assertNotIn("13800138000", version["base_content"])

    def test_updated_job_requires_a_fresh_analysis(self) -> None:
        update_job(
            self.job["id"],
            {"description": "负责全新的商业化产品方向，需要销售管理和渠道建设经验。"},
            self.db_path,
        )

        with self.assertRaisesRegex(ValueError, "重新生成岗位评估"):
            create_resume_version(self.job["id"], self.db_path)

    def test_baseline_version_uses_saved_resume_without_a_job(self) -> None:
        version = create_baseline_resume_version(self.db_path)

        self.assertIsNone(version["job_id"])
        self.assertIsNone(version["evaluation_id"])
        self.assertEqual(version["status"], "draft")
        self.assertEqual(version["template_id"], "classic")
        self.assertIn("测试用户", version["rendered_content"])
        self.assertIn("5年产品经验", version["rendered_content"])
        self.assertNotIn("13800138000", version["base_content"])
        self.assertNotIn("13800138000", version["rendered_content"])
        self.assertTrue(all(change["evidence"] for change in version["changes"]))
        self.assertEqual(list_resume_versions(db_path=self.db_path)[0]["id"], version["id"])
        self.assertEqual(list_resume_versions(self.job["id"], self.db_path), [])

        compact = update_resume_version(
            version["id"],
            template_id="compact",
            db_path=self.db_path,
        )
        self.assertEqual(compact["template_id"], "compact")

        docx_bytes, docx_name, _docx_type = export_resume_version(
            version["id"],
            "docx",
            self.db_path,
        )
        self.assertTrue(docx_name.endswith(".docx"))
        document = Document(BytesIO(docx_bytes))
        docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("5年产品经验", docx_text)
        self.assertNotIn("13800138000", docx_text)

    def test_baseline_version_does_not_require_job_evaluation(self) -> None:
        bare_job = create_job(
            {"job_title": "无分析岗位", "description": "任意描述，尚未生成评估。"},
            self.db_path,
        )
        with self.assertRaisesRegex(ValueError, "请先生成岗位决策"):
            create_resume_version(bare_job["id"], self.db_path)

        version = create_resume_version(db_path=self.db_path)
        self.assertIsNone(version["job_id"])
        self.assertIn("负责 Agent 产品规划", version["base_content"])

    def test_titled_capabilities_become_three_strength_entries(self) -> None:
        layout = split_resume_layout(
            "陈露鑫｜AI 应用工程师\n"
            "GitHub：https://github.com/example\n"
            "\n"
            "「AIGC 与大模型落地能力」：熟练掌握 LangChain。\n"
            "「AI 工程化全栈交付能力」：独立完成全栈交付。\n"
            "「产品从 0 到 1 落地迭代能力」：从需求到上线闭环。\n"
            "\n"
            "工作经历\n"
            "示例科技｜AI 应用工程师\n"
        )
        strengths = next(section for section in layout["main"] if section["kind"] == "strengths")
        summary_text = "\n".join(
            line
            for section in [*layout["sidebar"], *layout["main"]]
            if section["kind"] == "summary"
            for entry in section["entries"]
            for line in entry
        )
        self.assertEqual(strengths["label"], "个人优势")
        self.assertEqual(
            strengths["entries"],
            [
                ["「AIGC 与大模型落地能力」", "熟练掌握 LangChain。"],
                ["「AI 工程化全栈交付能力」", "独立完成全栈交付。"],
                ["「产品从 0 到 1 落地迭代能力」", "从需求到上线闭环。"],
            ],
        )
        self.assertNotIn("AIGC 与大模型落地能力", summary_text)
        self.assertEqual([section["kind"] for section in layout["main"] if section["kind"] == "experience"], ["experience"])

    def test_compact_layout_puts_skills_in_the_sidebar(self) -> None:
        layout = split_resume_layout(
            "陈露鑫｜后端工程师\n\n工作经历\n示例科技\n负责接口\n\n核心技能\nPython、FastAPI\n\n教育经历\n复旦大学"
        )
        self.assertEqual(layout["title"], "陈露鑫｜后端工程师")
        self.assertEqual([section["kind"] for section in layout["sidebar"]], ["skills", "education"])
        self.assertEqual([section["kind"] for section in layout["main"]], ["experience"])

    def test_compact_export_uses_a_two_column_table(self) -> None:
        with connect(self.db_path) as conn:
            conn.execute(
                "UPDATE profiles SET resume_text = ?, resume_redacted_text = ?",
                (
                    "测试用户｜后端\n\n工作经历\n示例科技\n负责接口\n\n核心技能\nPython、FastAPI\n\n教育经历\n复旦大学",
                    "测试用户｜后端\n\n工作经历\n示例科技\n负责接口\n\n核心技能\nPython、FastAPI\n\n教育经历\n复旦大学",
                ),
            )
        version = create_baseline_resume_version(self.db_path)
        update_resume_version(version["id"], template_id="compact", db_path=self.db_path)
        docx_bytes, _, _ = export_resume_version(version["id"], "docx", self.db_path)
        document = Document(BytesIO(docx_bytes))
        self.assertGreaterEqual(len(document.tables), 1)
        table_text = "\n".join(cell.text for cell in document.tables[0].rows[0].cells)
        self.assertIn("技能", table_text)
        self.assertIn("Python", table_text)
        self.assertIn("工作与实习经历", table_text)

        pdf_bytes, _, _ = export_resume_version(version["id"], "pdf", self.db_path)
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))
        pages = PdfReader(BytesIO(pdf_bytes)).pages
        text = "\n".join(page.extract_text() or "" for page in pages)
        self.assertIn("Python", text)
        self.assertIn("复旦大学", text)

    def test_export_allows_studio_certificate_headings(self) -> None:
        version = create_baseline_resume_version(self.db_path)
        body = next(change for change in version["changes"] if change["section_key"] == "body")
        update_resume_change(
            version["id"],
            body["id"],
            after_text="荣誉证书\n三好学生\n\n工作经历\n5年产品经验\n负责 Agent 产品规划。",
            db_path=self.db_path,
        )

        payload, filename, media_type = export_resume_version(
            version["id"], "docx", self.db_path
        )
        self.assertTrue(filename.endswith(".docx"))
        self.assertIn("wordprocessingml", media_type)
        self.assertGreater(len(payload), 100)

    def test_rendered_content_does_not_repeat_the_name_title(self) -> None:
        version = create_baseline_resume_version(self.db_path)
        body = next(change for change in version["changes"] if change["section_key"] == "body")
        update_resume_change(
            version["id"],
            body["id"],
            after_text=(
                "测试用户\n"
                "GitHub：https://github.com/example\n"
                "求职意向：AI应用工程师\n"
                "\n"
                "个人优势\n"
                "「AIGC 与大模型落地能力」：熟练掌握 LangChain 框架、RAG 检索增强，落地会议 AI\n"
                "问答，有效提升业务内容产出效\n"
                "率 60%+。\n"
                "\n"
                "工作经历\n"
                "示例科技｜AI 应用工程师 2025.07 - 2026.04\n"
                "\n"
                "项目经历\n"
                "智能会议总结（Summary）\n"
                "将新模型接入周期由\n"
                "3\n"
                "天缩短至 4h。\n"
                "单节\n"
                "点每日承载 3000+ 条任务。\n"
                "\n"
                "教育经历\n"
                "示例大学 软件工程专业 2021.09-2025.6\n"
                "\n"
                "荣誉证书\n"
                "英语：CET-6\n"
            ),
            db_path=self.db_path,
        )
        rendered = get_resume_version(version["id"], self.db_path)["rendered_content"]
        self.assertEqual(rendered.count("测试用户"), 1)
        self.assertIn("求职意向：AI应用工程师", rendered)
        self.assertIn("GitHub：https://github.com/example", rendered)

    def test_pdf_export_unwraps_wrapped_lines_and_styles_section_headings(self) -> None:
        version = create_baseline_resume_version(self.db_path)
        body = next(change for change in version["changes"] if change["section_key"] == "body")
        update_resume_change(
            version["id"],
            body["id"],
            after_text=(
                "# 测试用户\n"
                "\n"
                "测试用户\n"
                "GitHub：https://github.com/example\n"
                "求职意向：AI应用工程师\n"
                "\n"
                "个人优势\n"
                "「AIGC 与大模型落地能力」：使用 Python 与 FastAPI 完成内部原型，落地会议 AI\n"
                "问答，独立完成业务内容产出效\n"
                "率提升。\n"
                "\n"
                "工作经历\n"
                "示例科技｜AI 应用工程师\n"
                "负责 Agent 产品规划和需求分析。\n"
                "\n"
                "项目经历\n"
                "智能会议总结（Summary）\n"
                "使用 Python 与 FastAPI 完成内部原型，将接入周期由数\n"
                "天缩短，单节\n"
                "点每日承载会议任务。\n"
                "\n"
                "教育经历\n"
                "示例大学 软件工程专业\n"
                "\n"
                "荣誉证书\n"
                "英语：CET-6\n"
            ),
            db_path=self.db_path,
        )

        pdf_bytes, _, _ = export_resume_version(version["id"], "pdf", self.db_path)
        pages = PdfReader(BytesIO(pdf_bytes)).pages
        text = "\n".join(page.extract_text() or "" for page in pages)
        self.assertEqual(text.count("测试用户"), 1)
        self.assertIn("产出效率提升", text.replace("\n", ""))
        self.assertIn("由数天缩短", text.replace("\n", ""))
        self.assertIn("单节点每日承载", text.replace("\n", ""))
        self.assertIn("会议 AI 问答", text.replace("\n", ""))
        self.assertIn("荣誉证书", text)
        self.assertIn("CET-6", text)
        self.assertIn("github.com/example", text)
        self.assertNotIn("CareerLoop 定制简历", text)

        layout = split_resume_layout(get_resume_version(version["id"], self.db_path)["rendered_content"])
        self.assertEqual(layout["title"], "测试用户")
        self.assertTrue(any("github.com/example" in item for item in layout["contact"]))
        self.assertIn("AI应用工程师", layout["target"])
        self.assertEqual(
            [section["kind"] for section in layout["sections"]],
            ["strengths", "experience", "projects", "education", "honors"],
        )

    def test_compose_rendered_sections_drops_duplicate_hash_title(self) -> None:
        merged = compose_rendered_sections([
            "# 小程",
            "小程\n求职意向：AI应用工程师\n\n工作经历\n示例科技",
        ])
        self.assertEqual(merged.count("小程"), 1)
        self.assertIn("求职意向：AI应用工程师", merged)
        self.assertIn("工作经历", merged)

    def test_layout_drops_repeated_name_and_rejoins_broken_metric_lines(self) -> None:
        layout = split_resume_layout(
            "# 小程\n"
            "小程\n"
            "求职意向：AI应用工程师\n"
            "\n"
            "个人优势\n"
            "「AIGC 与大模型落地能力」：熟练掌握 LangChain。\n"
            "\n"
            "工作经历\n"
            "89Trillion | AI 应用工程师 2025.07 - 2026.04\n"
            "\n"
            "项目经历\n"
            "智能会议总结（Summary）\n"
            "将新模型接入周期由\n"
            "3\n"
            "天缩短至 4h。\n"
        )
        self.assertEqual(layout["title"], "小程")
        self.assertEqual(layout["target"], "AI应用工程师")
        body = "\n".join(
            line
            for section in layout["sections"]
            for entry in section["entries"]
            for line in entry
        )
        self.assertEqual(body.count("小程"), 0)
        self.assertIn("由3天缩短至4h", body.replace(" ", ""))

    def test_entry_heading_and_document_name_split_like_open_resume(self) -> None:
        self.assertEqual(
            split_entry_heading("89Trillion | AI 应用工程师 2025.07 - 2026.04"),
            ("89Trillion | AI 应用工程师", "2025.07 - 2026.04"),
        )
        self.assertEqual(
            split_entry_heading("2024.01 - 至今  星环科技｜前端工程师"),
            ("星环科技｜前端工程师", "2024.01 - 至今"),
        )
        self.assertEqual(
            split_entry_heading("CareerLoop 求职助手｜2025.03 - 至今"),
            ("CareerLoop 求职助手", "2025.03 - 至今"),
        )
        self.assertEqual(split_document_name("陈露鑫｜AI 应用工程师"), ("陈露鑫", "AI 应用工程师"))
        self.assertEqual(split_document_name("CareerLoop 求职助手｜2025.03 - 至今"), ("CareerLoop 求职助手｜2025.03 - 至今", ""))
        self.assertEqual(contact_link_target("GitHub：https://github.com/example"), "https://github.com/example")
        self.assertEqual(contact_link_target("name@example.com"), "mailto:name@example.com")

    def test_export_uses_selected_style_colors(self) -> None:
        version = create_baseline_resume_version(self.db_path)
        update_resume_version(version["id"], style_id="wine", db_path=self.db_path)
        docx_bytes, _, _ = export_resume_version(version["id"], "docx", self.db_path)
        document = Document(BytesIO(docx_bytes))
        self.assertEqual(str(document.styles["Title"].font.color.rgb), "5C2433")
        self.assertEqual(str(document.styles["Heading 2"].font.color.rgb), "8B3D52")
        headings = [paragraph for paragraph in document.paragraphs if paragraph.style.name == "Heading 1"]
        self.assertTrue(headings)
        heading_xml = headings[0]._p.xml
        self.assertIn("w:bottom", heading_xml)
        self.assertNotIn("CareerLoop 定制简历", "\n".join(paragraph.text for paragraph in document.paragraphs))


if __name__ == "__main__":
    unittest.main()
