from __future__ import annotations

import re
from html import escape
from io import BytesIO
from pathlib import Path
from typing import Any

from .layout import (
    contact_link_target,
    skill_tags,
    split_document_name,
    split_entry_heading,
    split_resume_layout,
)


def build_docx(
    title: str,
    content: str,
    template_id: str,
    palette: dict[str, str | bool],
    scale: float,
) -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    compact = template_id == "compact"
    minimal = template_id == "minimal"
    section.top_margin = Cm((1.15 if compact else 1.9 if minimal else 1.6) * scale)
    section.bottom_margin = Cm((1.15 if compact else 1.7 if minimal else 1.5) * scale)
    section.left_margin = Cm((1.35 if compact else 1.9 if minimal else 1.7) * scale)
    section.right_margin = Cm((1.35 if compact else 1.9 if minimal else 1.7) * scale)
    section.start_type = WD_SECTION.NEW_PAGE

    font_name = _document_font_name(serif=bool(palette["serif"]))
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(max(8.5, (9.4 if compact else 10.4) * scale))
    normal.font.color.rgb = RGBColor.from_string(str(palette["body"]))
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    for style_name, size, color in (
        ("Title", (17 if compact else 21) * scale, str(palette["title"])),
        ("Heading 1", (11 if compact else 12.5) * scale, str(palette["heading"])),
        ("Heading 2", (10 if compact else 11.5) * scale, str(palette["heading"])),
        ("List Bullet", (8.6 if compact else 10.2) * scale, str(palette["body"])),
    ):
        style = styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    if not (compact and _append_compact_docx_columns(document, content, palette, scale)):
        _append_linear_docx(document, content, palette, scale)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_run = footer.add_run()
    page_run.font.name = font_name
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = RGBColor.from_string("9AA3A8")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_separator = OxmlElement("w:fldChar")
    field_separator.set(qn("w:fldCharType"), "separate")
    field_value = OxmlElement("w:t")
    field_value.text = "1"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    page_run._r.extend(
        (field_begin, instruction, field_separator, field_value, field_end)
    )
    document.core_properties.title = title
    document.core_properties.subject = "岗位定制简历"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf(
    title: str,
    content: str,
    template_id: str,
    palette: dict[str, str | bool],
    scale: float,
    layout: dict[str, Any] | None = None,
) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import KeepInFrame, SimpleDocTemplate
    except ImportError as exc:
        raise RuntimeError("PDF 导出依赖未安装，请安装 reportlab") from exc

    fonts = _pdf_fonts(serif=bool(palette["serif"]))
    compact = template_id == "compact"
    minimal = template_id == "minimal"
    one_page = bool((layout or {}).get("one_page"))
    output = BytesIO()
    side = (12 if compact else 16 if minimal else 15) * mm * scale
    top = (11 if compact else 15 if minimal else 14) * mm * scale
    bottom = (12 if compact else 14 if minimal else 13) * mm * scale
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=side,
        leftMargin=side,
        topMargin=top,
        bottomMargin=bottom,
        title=title,
        author="CareerLoop",
    )
    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "ResumeBody",
        parent=base["BodyText"],
        fontName=fonts["regular"],
        fontSize=(8.2 if compact else 9.6) * scale,
        leading=(11.4 if compact else 13.6) * scale,
        textColor=colors.HexColor(f"#{palette['body']}"),
        spaceAfter=(1.5 if compact else 2.5) * scale,
        wordWrap="CJK",
    )
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=body,
        fontName=fonts["bold"],
        fontSize=(16 if compact else 20) * scale,
        leading=(20 if compact else 24) * scale,
        textColor=colors.HexColor(f"#{palette['title']}"),
        spaceAfter=1 * scale,
        wordWrap="CJK",
    )
    role_style = ParagraphStyle(
        "ResumeRole",
        parent=body,
        fontName=fonts["bold"],
        fontSize=(9.2 if compact else 10.4) * scale,
        leading=(12 if compact else 13.5) * scale,
        textColor=colors.HexColor(f"#{palette['heading']}"),
        spaceAfter=2 * scale,
        wordWrap="CJK",
    )
    meta = ParagraphStyle(
        "ResumeMeta",
        parent=body,
        fontSize=(7.8 if compact else 8.6) * scale,
        leading=(10.6 if compact else 12) * scale,
        textColor=colors.HexColor(f"#{palette['accent']}"),
        spaceAfter=1.5 * scale,
        wordWrap="CJK",
    )
    heading = ParagraphStyle(
        "ResumeHeading",
        parent=body,
        fontName=fonts["bold"],
        fontSize=(10.2 if compact else 11.2) * scale,
        leading=(13 if compact else 14.5) * scale,
        textColor=colors.HexColor(f"#{palette['heading']}"),
        spaceBefore=(6 if compact else 7) * scale,
        spaceAfter=1 * scale,
        wordWrap="CJK",
        keepWithNext=True,
    )
    entry_title = ParagraphStyle(
        "ResumeEntryTitle",
        parent=body,
        fontName=fonts["bold"],
        textColor=colors.HexColor(f"#{palette['title']}"),
        spaceAfter=1 * scale,
        keepWithNext=True,
        wordWrap="CJK",
    )
    entry_date = ParagraphStyle(
        "ResumeEntryDate",
        parent=meta,
        alignment=2,
        spaceAfter=1 * scale,
        keepWithNext=True,
        wordWrap="CJK",
    )
    bullet = ParagraphStyle(
        "ResumeBullet",
        parent=body,
        leftIndent=11,
        firstLineIndent=-7,
        bulletIndent=2,
        spaceAfter=(1.2 if compact else 1.8) * scale,
        wordWrap="CJK",
    )
    chip = ParagraphStyle(
        "ResumeChip",
        parent=body,
        fontName=fonts["bold"],
        fontSize=(7.4 if compact else 8) * scale,
        leading=(10 if compact else 11) * scale,
        textColor=colors.HexColor(f"#{palette['heading']}"),
        alignment=0,
        wordWrap="",
        splitLongWords=0,
    )
    styles = {
        "title": title_style,
        "role": role_style,
        "heading": heading,
        "body": body,
        "bullet": bullet,
        "meta": meta,
        "entry_title": entry_title,
        "entry_date": entry_date,
        "chip": chip,
    }
    story = _pdf_story(content, styles, palette, compact=compact, scale=scale)
    if one_page and story:
        frame_width = A4[0] - side * 2
        frame_height = A4[1] - top - bottom - 6 * mm
        story = [KeepInFrame(frame_width, frame_height, story, mode="shrink")]

    accent = colors.HexColor(f"#{palette['heading']}")

    def draw_chrome(canvas, doc, *, first: bool) -> None:
        canvas.saveState()
        if first:
            canvas.setFillColor(accent)
            canvas.rect(0, A4[1] - 3.2 * mm, A4[0], 3.2 * mm, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor("#9AA3A8"))
        canvas.setFont(fonts["regular"], 8)
        canvas.drawCentredString(A4[0] / 2, 8 * mm, str(doc.page))
        canvas.restoreState()

    document.build(
        story,
        onFirstPage=lambda canvas, doc: draw_chrome(canvas, doc, first=True),
        onLaterPages=lambda canvas, doc: draw_chrome(canvas, doc, first=False),
    )
    return output.getvalue()


def _append_linear_docx(document, content: str, palette: dict[str, str | bool], scale: float) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    layout = split_resume_layout(content)
    if layout["title"] or layout.get("contact") or layout.get("target"):
        _append_docx_header(document, layout, palette, scale)
    sections = layout.get("sections") or [*layout["sidebar"], *layout["main"]]
    if not sections and not layout["title"]:
        _append_docx_plain_lines(document, content, palette)
        return
    for section in sections:
        heading = document.add_heading(str(section["label"]), level=1)
        heading.paragraph_format.space_before = Pt(10 * scale)
        heading.paragraph_format.space_after = Pt(3 * scale)
        heading.paragraph_format.keep_with_next = True
        _set_paragraph_bottom_border(heading, str(palette["rule"]))
        _append_docx_section_body(document, section, palette, scale)


def _append_docx_plain_lines(document, content: str, palette: dict[str, str | bool]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    for line in content.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith("# "):
            paragraph = document.add_paragraph(text[2:], style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text.startswith("## "):
            heading = document.add_heading(text[3:], level=1)
            _set_paragraph_bottom_border(heading, str(palette["rule"]))
        elif text.startswith("- "):
            paragraph = document.add_paragraph(text[2:], style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
        else:
            paragraph = document.add_paragraph(text)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.15


def _append_compact_docx_columns(
    document,
    content: str,
    palette: dict[str, str | bool],
    scale: float,
) -> bool:
    from docx.shared import Cm

    layout = split_resume_layout(content)
    if not layout["sidebar"] or not layout["main"]:
        return False
    if layout["title"] or layout.get("contact") or layout.get("target"):
        _append_docx_header(document, layout, palette, scale)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(5.6)
    table.columns[1].width = Cm(11.4)
    _set_cell_right_border(table.cell(0, 0), str(palette["rule"]))
    _fill_docx_layout_cell(table.cell(0, 0), layout["sidebar"], palette, scale)
    _fill_docx_layout_cell(table.cell(0, 1), layout["main"], palette, scale)
    return True


def _append_docx_header(
    document,
    layout: dict[str, Any],
    palette: dict[str, str | bool],
    scale: float,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    name, role = split_document_name(str(layout.get("title") or ""))
    if name:
        paragraph = document.add_paragraph(name, style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2 * scale)
    if role and role != name:
        paragraph = document.add_paragraph(role)
        paragraph.paragraph_format.space_after = Pt(2 * scale)
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(str(palette["heading"]))
            run.font.size = Pt(11 * scale)
    _append_docx_header_meta(document, layout, palette, scale)


def _append_docx_header_meta(
    document,
    layout: dict[str, Any],
    palette: dict[str, str | bool],
    scale: float,
) -> None:
    from docx.shared import Pt, RGBColor

    accent = RGBColor.from_string(str(palette["accent"]))
    contact = [str(item).strip() for item in layout.get("contact") or [] if str(item).strip()]
    if contact:
        paragraph = document.add_paragraph("  ·  ".join(contact))
        paragraph.paragraph_format.space_after = Pt(1 * scale)
        for run in paragraph.runs:
            run.font.color.rgb = accent
            run.font.size = Pt(9.5 * scale)
    target = str(layout.get("target") or "").strip()
    if target:
        label = target if target.startswith("求职") else f"求职意向：{target}"
        paragraph = document.add_paragraph(label)
        paragraph.paragraph_format.space_after = Pt(2 * scale)
        for run in paragraph.runs:
            run.font.color.rgb = accent
            run.font.size = Pt(9.5 * scale)
    if contact or layout.get("title") or target:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6 * scale)
        _set_paragraph_bottom_border(spacer, str(palette["rule"]))


def _append_docx_section_body(
    document,
    section: dict[str, Any],
    palette: dict[str, str | bool],
    scale: float,
) -> None:
    if section["kind"] == "skills":
        _append_docx_skill_chips(document, skill_tags(section["entries"]), palette, scale)
        return
    for entry in section["entries"]:
        _append_docx_entry(document, entry, palette, scale)


def _append_docx_entry(
    document,
    entry: list[str],
    palette: dict[str, str | bool],
    scale: float,
) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Pt, RGBColor

    if not entry:
        return
    title, date = split_entry_heading(entry[0])
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2 * scale)
    paragraph.paragraph_format.keep_with_next = len(entry) > 1
    title_run = paragraph.add_run(title or entry[0])
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(str(palette["title"]))
    if date:
        usable = _docx_usable_width(document)
        paragraph.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
        paragraph.add_run("\t")
        date_run = paragraph.add_run(date)
        date_run.font.color.rgb = RGBColor.from_string(str(palette["accent"]))
        date_run.font.size = Pt(9.2 * scale)
    for line in entry[1:]:
        if hasattr(document, "_tc"):
            bullet = document.add_paragraph(f"• {line}")
        else:
            bullet = document.add_paragraph(line, style="List Bullet")
        bullet.paragraph_format.space_after = Pt(1.5 * scale)


def _append_docx_skill_chips(
    document,
    tags: list[str],
    palette: dict[str, str | bool],
    scale: float,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    if not tags:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6 * scale)
    heading = RGBColor.from_string(str(palette["heading"]))
    chip = str(palette.get("chip") or "EEF3F8")
    for index, tag in enumerate(tags):
        if index:
            spacer = paragraph.add_run("  ")
            spacer.font.size = Pt(8 * scale)
        run = paragraph.add_run(f"  {tag}  ")
        run.font.size = Pt(8.4 * scale)
        run.font.color.rgb = heading
        run.bold = True
        rpr = run._element.get_or_add_rPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), chip)
        shading.set(qn("w:val"), "clear")
        rpr.append(shading)


def _fill_docx_layout_cell(
    cell,
    sections: list[dict],
    palette: dict[str, str | bool],
    scale: float,
) -> None:
    from docx.shared import Pt, RGBColor

    cell.text = ""
    first = True
    heading_color = RGBColor.from_string(str(palette["heading"]))
    for section in sections:
        heading = cell.paragraphs[0] if first else cell.add_paragraph()
        heading.text = str(section["label"])
        for run in heading.runs:
            run.bold = True
            run.font.color.rgb = heading_color
        heading.paragraph_format.space_after = Pt(3 * scale)
        heading.paragraph_format.keep_with_next = True
        _set_paragraph_bottom_border(heading, str(palette["rule"]))
        first = False
        if section["kind"] == "skills":
            _append_docx_skill_chips(cell, skill_tags(section["entries"]), palette, scale)
            continue
        for entry in section["entries"]:
            _append_docx_entry(cell, entry, palette, scale)


def _docx_usable_width(container) -> int:
    from docx.shared import Cm

    width = getattr(container, "width", None)
    if width:
        return int(width)
    try:
        section = container.sections[0]
        return int(section.page_width - section.left_margin - section.right_margin)
    except AttributeError:
        return int(Cm(11.4))


def _set_paragraph_bottom_border(paragraph, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph_props = paragraph._p.get_or_add_pPr()
    borders = paragraph_props.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_props.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    existing = borders.find(qn("w:bottom"))
    if existing is not None:
        borders.remove(existing)
    borders.append(bottom)


def _set_cell_right_border(cell, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    props = cell._tc.get_or_add_tcPr()
    borders = props.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        props.append(borders)
    right = OxmlElement("w:right")
    right.set(qn("w:val"), "single")
    right.set(qn("w:sz"), "6")
    right.set(qn("w:color"), color)
    existing = borders.find(qn("w:right"))
    if existing is not None:
        borders.remove(existing)
    borders.append(right)


def _pdf_story(
    content: str,
    styles: dict,
    palette: dict[str, str | bool],
    *,
    compact: bool,
    scale: float,
):
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import Table, TableStyle

    layout = split_resume_layout(content)
    story = []
    header = _pdf_header_flowables(layout, styles, palette, scale=scale)
    if header:
        story.extend(header)
    sections = layout.get("sections") or [*layout["sidebar"], *layout["main"]]
    if compact and layout["sidebar"] and layout["main"]:
        sidebar = _pdf_section_flowables(
            layout["sidebar"],
            styles,
            palette,
            scale=scale,
            grouped=False,
            narrow=True,
            max_width=54 * mm,
        )
        main = _pdf_section_flowables(
            layout["main"],
            styles,
            palette,
            scale=scale,
            grouped=False,
            max_width=114 * mm,
        )
        table = Table([[sidebar, main]], colWidths=[58 * mm, 118 * mm])
        table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (0, 0), 0),
            ("RIGHTPADDING", (0, 0), (0, 0), 8),
            ("LEFTPADDING", (1, 0), (1, 0), 12),
            ("RIGHTPADDING", (1, 0), (1, 0), 0),
            ("LINEAFTER", (0, 0), (0, 0), 0.5, colors.HexColor(f"#{palette['rule']}")),
        ]))
        story.append(table)
        return story
    if not sections:
        return story or _pdf_plain_story(content, styles)
    for section in sections:
        story.extend(_pdf_section_flowables([section], styles, palette, scale=scale, max_width=170 * mm))
    return story


def _pdf_header_flowables(layout: dict[str, Any], styles: dict, palette: dict[str, str | bool], *, scale: float):
    from reportlab.lib import colors
    from reportlab.platypus import HRFlowable, Paragraph, Spacer

    flowables = []
    name, role = split_document_name(str(layout.get("title") or ""))
    if name:
        flowables.append(Paragraph(_pdf_text(name), styles["title"]))
    if role and role != name:
        flowables.append(Paragraph(_pdf_text(role), styles["role"]))
    contact = [
        str(item).strip()
        for item in layout.get("contact") or []
        if str(item).strip() and str(item).strip() != name
    ]
    if contact:
        flowables.append(Paragraph(_pdf_contact_line(contact, str(palette["accent"])), styles["meta"]))
    target = str(layout.get("target") or "").strip()
    if target:
        label = target if target.startswith("求职") else f"求职意向：{target}"
        flowables.append(Paragraph(_pdf_text(label), styles["meta"]))
    if flowables:
        flowables.append(Spacer(1, 3 * scale))
        flowables.append(HRFlowable(
            width="100%",
            thickness=0.6,
            color=colors.HexColor(f"#{palette['rule']}"),
            spaceBefore=0,
            spaceAfter=4 * scale,
        ))
    return flowables


def _pdf_plain_story(content: str, styles: dict):
    from reportlab.platypus import Paragraph, Spacer

    story = []
    for line in content.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith("# "):
            story.append(Paragraph(_pdf_text(text[2:]), styles["title"]))
        elif text.startswith("## "):
            story.append(Paragraph(_pdf_text(text[3:]), styles["heading"]))
        elif text.startswith("- "):
            story.append(Paragraph(f'<font color="#3E8E6B">•</font> {_pdf_text(text[2:])}', styles["bullet"]))
        else:
            story.append(Paragraph(_pdf_text(text), styles["body"]))
    if not story:
        story.append(Spacer(1, 1))
    return story


def _pdf_section_flowables(
    sections: list[dict],
    styles: dict,
    palette: dict[str, str | bool],
    *,
    scale: float,
    grouped: bool = True,
    narrow: bool = False,
    max_width: float | None = None,
):
    from reportlab.lib import colors
    from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, Spacer

    flowables = []
    accent = str(palette["accent"])
    for section in sections:
        heading_block = [
            Paragraph(_pdf_text(str(section["label"])), styles["heading"]),
            HRFlowable(
                width="100%",
                thickness=0.7,
                color=colors.HexColor(f"#{palette['rule']}"),
                spaceBefore=0,
                spaceAfter=4 * scale,
            ),
        ]
        if section["kind"] == "skills":
            chips = _pdf_skill_chips(
                skill_tags(section["entries"]),
                styles,
                palette,
                scale=scale,
                max_width=max_width,
            )
            flowables.extend(heading_block)
            flowables.extend(chips)
            continue
        for entry in section["entries"]:
            block = [
                *heading_block,
                *_pdf_entry_flowables(entry, styles, accent, scale=scale, narrow=narrow),
            ]
            if grouped:
                flowables.append(KeepTogether(block))
            else:
                flowables.extend(block)
            heading_block = []
        if heading_block:
            flowables.extend(heading_block)
    return flowables


def _pdf_entry_flowables(
    entry: list[str],
    styles: dict,
    accent: str,
    *,
    scale: float,
    narrow: bool = False,
):
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    if not entry:
        return []
    title, date = split_entry_heading(entry[0])
    title_para = Paragraph(_pdf_text(title or entry[0]), styles["entry_title"])
    if date:
        header = Table(
            [[title_para, Paragraph(_pdf_text(date), styles["entry_date"])]],
            colWidths=[None, (22 if narrow else 36) * mm],
        )
        header.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ]))
        flowables = [header]
    else:
        flowables = [title_para]
    for line in entry[1:]:
        flowables.append(
            Paragraph(f'<font color="#{accent}">•</font> {_pdf_text(line)}', styles["bullet"])
        )
    flowables.append(Spacer(1, 2 * scale))
    return flowables


def _pdf_skill_chips(
    tags: list[str],
    styles: dict,
    palette: dict[str, str | bool],
    *,
    scale: float,
    max_width: float | None = None,
):
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    if not tags:
        return []
    font_name = styles["chip"].fontName
    font_size = styles["chip"].fontSize
    pad = 14 * scale
    gap = 5 * scale
    limit = max_width or 170 * mm
    chip_bg = colors.HexColor(f"#{palette.get('chip') or 'EEF3F8'}")
    chip_style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), chip_bg),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 2.4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2.4),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
    ])

    if max_width and max_width < 70 * mm:
        data = [[Paragraph(f"<nobr>{_pdf_text(tag)}</nobr>", styles["chip"])] for tag in tags]
        table = Table(data, colWidths=[limit], hAlign="LEFT")
        table.setStyle(chip_style)
        table.setStyle(TableStyle([
            ("TOPPADDING", (0, 1), (-1, -1), 3.5),
            ("BOTTOMPADDING", (0, 0), (-1, -2), 3.5),
        ]))
        return [table, Spacer(1, 3 * scale)]

    rows: list[list[tuple[str, float]]] = []
    current: list[tuple[str, float]] = []
    current_width = 0.0
    for tag in tags:
        width = min(limit, stringWidth(tag, font_name, font_size) + pad)
        if current and current_width + gap + width > limit:
            rows.append(current)
            current = [(tag, width)]
            current_width = width
        else:
            current.append((tag, width))
            current_width += width if len(current) == 1 else gap + width
    if current:
        rows.append(current)

    flowables = []
    for row in rows:
        cells = [Paragraph(f"<nobr>{_pdf_text(tag)}</nobr>", styles["chip"]) for tag, _width in row]
        widths = [width for _tag, width in row]
        table = Table([cells], colWidths=widths, hAlign="LEFT")
        table.setStyle(chip_style)
        table.setStyle(TableStyle([
            ("RIGHTPADDING", (0, 0), (-2, -1), 8),
        ]))
        flowables.append(table)
    flowables.append(Spacer(1, 3 * scale))
    return flowables


def _pdf_contact_line(items: list[str], accent: str) -> str:
    parts = []
    for item in items:
        target = contact_link_target(item)
        text = _pdf_text(item)
        if target:
            parts.append(f'<link href="{escape(target, quote=True)}"><font color="#{accent}">{text}</font></link>')
        else:
            parts.append(text)
    return "  ·  ".join(parts)


_PDF_METRIC = re.compile(
    r"(?:≤|≥|>=|<=)?\d+(?:\.\d+)?\s*(?:%|h|min|s|ms|倍)?\+?"
)


def _pdf_text(value: str) -> str:
    escaped = escape(value)
    return _PDF_METRIC.sub(lambda match: f"<nobr>{match.group(0)}</nobr>", escaped)


def _document_font_name(*, serif: bool = False) -> str:
    if serif:
        if Path("/System/Library/Fonts/Supplemental/Songti.ttc").exists():
            return "Songti SC"
        if Path("/Library/Fonts/Songti.ttc").exists():
            return "Songti SC"
        if Path("C:/Windows/Fonts/simsun.ttc").exists():
            return "SimSun"
        return "Times New Roman"
    if Path("/System/Library/Fonts/PingFang.ttc").exists():
        return "PingFang SC"
    if Path("/System/Library/Fonts/Hiragino Sans GB.ttc").exists():
        return "Hiragino Sans GB"
    if Path("/Library/Fonts/Arial Unicode.ttf").exists():
        return "Arial Unicode MS"
    if Path("C:/Windows/Fonts/msyh.ttc").exists():
        return "Microsoft YaHei"
    return "Noto Sans CJK SC"


def _pdf_fonts(*, serif: bool) -> dict[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.pdfbase.ttfonts import TTFont

    regular_name = "CareerLoopResume"
    bold_name = "CareerLoopResumeBold"
    for path, regular_index, bold_index in _pdf_font_candidates(serif=serif):
        if not path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(regular_name, str(path), subfontIndex=regular_index))
            if stringWidth("中", regular_name, 10) <= 0:
                continue
            try:
                pdfmetrics.registerFont(TTFont(bold_name, str(path), subfontIndex=bold_index))
                if stringWidth("中", bold_name, 10) <= 0:
                    bold_name = regular_name
            except Exception:
                bold_name = regular_name
            return {"regular": regular_name, "bold": bold_name}
        except Exception:
            continue
    fallback = "STSong-Light"
    pdfmetrics.registerFont(UnicodeCIDFont(fallback))
    return {"regular": fallback, "bold": fallback}


def _pdf_font_candidates(*, serif: bool) -> tuple[tuple[Path, int, int], ...]:
    if serif:
        return (
            (Path("/System/Library/Fonts/Supplemental/Songti.ttc"), 0, 3),
            (Path("/Library/Fonts/Songti.ttc"), 0, 3),
            (Path("C:/Windows/Fonts/simsun.ttc"), 0, 0),
        )
    return (
        (Path("/System/Library/Fonts/PingFang.ttc"), 2, 8),
        (Path("/System/Library/Fonts/Hiragino Sans GB.ttc"), 0, 1),
        (Path("/Library/Fonts/Arial Unicode.ttf"), 0, 0),
        (Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf"), 0, 0),
        (Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"), 0, 0),
        (Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"), 0, 0),
        (Path("C:/Windows/Fonts/msyh.ttf"), 0, 0),
        (Path("C:/Windows/Fonts/msyh.ttc"), 0, 0),
        (Path("C:/Windows/Fonts/simhei.ttf"), 0, 0),
    )
