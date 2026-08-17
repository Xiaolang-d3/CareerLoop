from __future__ import annotations

from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from dataclasses import dataclass, field
import re
import unicodedata

from docx import Document
from pypdf import PdfReader


MAX_RESUME_BYTES = 8 * 1024 * 1024
SUPPORTED_RESUME_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
_SECTION_HEADING = re.compile(
    r"^(个人优势|核心优势|个人亮点|能力特长|核心竞争力|个人简介|教育经历|工作经历|项目经历|实习经历|校园经历|"
    r"专业技能|技能特长|自我评价|荣誉奖项|求职意向|基本信息|联系方式|"
    r"论文专利|所获荣誉|工作经验|项目经验|个人经历)[:：]?$"
)
_STRUCTURAL_LINE = re.compile(
    r"^([-–—*•●▪◦·]\s*|\d{1,2}[\.、\)]\s+|\d{4}\s*[\./年\-])"
)
_FIELD_LINE = re.compile(r"^.{1,16}[:：]")


@dataclass
class ResumeParseResult:
    text: str
    parser: str
    warnings: list[str] = field(default_factory=list)


def parse_resume(filename: str, content: bytes) -> str:
    return parse_resume_result(filename, content).text


def parse_resume_result(filename: str, content: bytes, mode: str = "fast") -> ResumeParseResult:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_RESUME_SUFFIXES:
        raise ValueError("仅支持 PDF、Word（.docx）和文本简历")
    if not content:
        raise ValueError("简历文件为空")
    if len(content) > MAX_RESUME_BYTES:
        raise ValueError("简历文件不能超过 8MB")

    if mode not in {"fast", "enhanced"}:
        raise ValueError("解析模式必须是 fast 或 enhanced")

    warnings: list[str] = []
    if mode == "enhanced" and suffix in {".pdf", ".docx"}:
        try:
            text = _parse_with_docling(filename, content)
            parser = "docling"
        except Exception:
            text = _parse_lightweight(suffix, content)
            parser = "lightweight"
            warnings.append("增强解析暂不可用，已自动使用快速解析")
    else:
        text = _parse_lightweight(suffix, content)
        parser = "lightweight"

    normalized = normalize_resume_text(text)
    if len(normalized) < 20:
        raise ValueError("未能从简历中提取足够文字；扫描版 PDF 请使用增强解析")
    return ResumeParseResult(normalized[:100_000], parser, warnings)


def normalize_resume_text(text: str) -> str:
    """Keep resume text readable after PDF/Word font extraction.

    Some templates encode bullet points in the Unicode private-use area (for
    example ``\uf0b7``). Those glyphs are not meaningful text and otherwise end
    up in the editable resume. Convert line-leading private-use glyphs into a
    normal dash and discard other invisible/control characters without touching
    URLs, Chinese text, punctuation, or tables.

    PDF/Word extractors also insert visual line breaks mid-sentence. Join those
    wrapped fragments locally; do not send the resume to a model for cleanup.
    """
    normalized = unicodedata.normalize("NFC", text).replace("\r\n", "\n").replace("\r", "\n")
    output: list[str] = []
    line_start = True
    for character in normalized:
        category = unicodedata.category(character)
        if character == "\n":
            output.append(character)
            line_start = True
            continue
        if character == "\t":
            output.append(character)
            line_start = False
            continue
        if category == "Co":
            if line_start:
                output.append("-")
            continue
        if category == "Cf" or category == "Cc":
            continue
        if character in {"•", "●", "▪", "◦", "■", "◆"} and line_start:
            output.append("-")
            continue
        output.append(character)
        line_start = character.isspace()

    cleaned = [
        re.sub(r"^-\s*", "- ", line.strip())
        for line in "".join(output).splitlines()
        if line.strip()
    ]
    return "\n".join(_split_jammed_profile_lines(_unwrap_extracted_lines(cleaned)))


_PROFILE_FIELD_LABELS = (
    "电话|手机|邮箱|邮件|微信|地址|住址|联系方式|GitHub|Github|LinkedIn|"
    "求职意向|意向岗位|目标职位|求职目标|英语|日语|普通话|语言"
)
_PROFILE_FIELD_PATTERN = re.compile(rf"(?:{_PROFILE_FIELD_LABELS})[:：]", re.I)
_JAMMED_CERT_GITHUB = re.compile(r"(CET-?\d)(?=[A-Za-z\u4e00-\u9fff])", re.I)
_JAMMED_GITHUB_LABEL = re.compile(r"([^\s|/｜])((?:GitHub|Github|LinkedIn)[:：])", re.I)


def _split_jammed_profile_line(line: str) -> list[str]:
    spaced = _JAMMED_CERT_GITHUB.sub(r"\1 ", line)
    spaced = _JAMMED_GITHUB_LABEL.sub(r"\1 \2", spaced)
    starts = [match.start() for match in _PROFILE_FIELD_PATTERN.finditer(spaced)]
    if len(starts) < 2:
        return [spaced.strip() or line]
    chunks: list[str] = []
    if starts[0] > 0:
        prefix = spaced[: starts[0]].strip()
        if prefix:
            chunks.append(prefix)
    for index, start in enumerate(starts):
        end = starts[index + 1] if index + 1 < len(starts) else len(spaced)
        chunk = spaced[start:end].strip()
        if chunk:
            chunks.append(chunk)
    return chunks or [line]


def _split_jammed_profile_lines(lines: list[str]) -> list[str]:
    expanded: list[str] = []
    for line in lines:
        expanded.extend(_split_jammed_profile_line(line) if line else [line])
    return expanded


def _is_cjk(character: str) -> bool:
    return "\u4e00" <= character <= "\u9fff"


def _is_heading(line: str) -> bool:
    return bool(_SECTION_HEADING.match(line.strip()))


def _is_structural(line: str) -> bool:
    stripped = line.strip()
    return bool(_is_heading(stripped) or _STRUCTURAL_LINE.match(stripped))


def _should_join_extracted_lines(previous: str, current: str) -> bool:
    if not previous or not current:
        return False
    if _is_heading(previous) or _is_structural(current):
        return False
    if _FIELD_LINE.match(previous) or _FIELD_LINE.match(current):
        return False
    if previous.endswith(("，", "、", ",", "；")):
        return True
    if previous.endswith(("。", "！", "？", "：", ":", ";", "!", "?")):
        return False
    if len(previous) < 16:
        return False
    return True


def _join_extracted_lines(previous: str, current: str) -> str:
    left, right = previous[-1], current[0]
    if previous.endswith(","):
        return f"{previous} {current}" if not current[0].isspace() else previous + current
    if _is_cjk(left) and _is_cjk(right):
        return previous + current
    if left.isalnum() and right.isalnum():
        return f"{previous} {current}"
    return previous + current


def _unwrap_extracted_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    for line in lines:
        if merged and _should_join_extracted_lines(merged[-1], line):
            merged[-1] = _join_extracted_lines(merged[-1], line)
        else:
            merged.append(line)
    return merged


def _parse_lightweight(suffix: str, content: bytes) -> str:
    if suffix == ".pdf":
        return "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)
    if suffix == ".docx":
        document = Document(BytesIO(content))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            blocks.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(blocks)
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError:
        return content.decode("gb18030")


def _parse_with_docling(filename: str, content: bytes) -> str:
    # Docling is intentionally imported only for an explicit enhanced parse. It
    # is large and may need to initialize local OCR models on first use.
    from docling.document_converter import DocumentConverter

    with TemporaryDirectory(prefix="careerloop-resume-") as directory:
        path = Path(directory) / Path(filename).name
        path.write_bytes(content)
        result = DocumentConverter().convert(path)
        return result.document.export_to_markdown()
