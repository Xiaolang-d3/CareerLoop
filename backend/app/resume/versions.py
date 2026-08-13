from __future__ import annotations

from html import escape
from io import BytesIO
from pathlib import Path
from typing import Any, Literal

from ..db import connect, json_dump, row_to_dict
from ..jobs.evaluations import get_latest_completed_job_evaluation
from ..candidate_core import get_candidate_context, verify_candidate_material
from ..profile_intelligence import extract_skills


ResumeDecision = Literal["pending", "accepted", "rejected"]
ResumeVersionStatus = Literal["draft", "final"]
ResumeExportFormat = Literal["docx", "pdf"]
ResumeTemplate = Literal["classic", "compact", "minimal"]

_RESUME_TEMPLATES = {"classic", "compact", "minimal"}


def create_resume_version(
    job_id: int,
    db_path: str | Path | None = None,
) -> dict[str, Any]:
    with connect(db_path) as conn:
        job_row = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
        if job_row is None:
            raise ValueError("岗位项目不存在")
        profile_row = conn.execute(
            "SELECT * FROM profiles ORDER BY updated_at DESC, id DESC LIMIT 1"
        ).fetchone()
        if profile_row is None:
            raise ValueError("请先保存人物画像和简历")
        version_number = conn.execute(
            "SELECT COUNT(*) AS count FROM resume_versions WHERE job_id = ?",
            (job_id,),
        ).fetchone()["count"] + 1

    evaluation = get_latest_completed_job_evaluation(job_id, db_path=db_path)
    if evaluation is None:
        raise ValueError("请先生成岗位决策与证据报告")
    if evaluation["status"] not in {"completed", "partial_failed"}:
        raise ValueError("岗位决策报告尚未完成")
    if evaluation.get("is_stale"):
        raise ValueError("岗位、候选人知识或职业策略已更新，请重新生成岗位评估")

    job = row_to_dict(job_row) or {}
    profile = row_to_dict(profile_row) or {}
    resume_context = get_candidate_context(
        "resume", profile_id=int(profile["id"]),
        strategy_id=evaluation.get("strategy_id"), db_path=db_path,
    )
    resume_text = "\n".join(item["statement"] for item in resume_context["confirmed_facts"])
    if not resume_text.strip():
        raise ValueError("当前隐私模式下没有可用简历文本")

    title_parts = [
        str(job.get("company_name") or "").strip(),
        str(job.get("job_title") or "").strip(),
    ]
    title = " · ".join(part for part in title_parts if part)
    title = f"{title or '岗位'}定制简历 V{version_number}"
    material_view = _evaluation_material_view(evaluation, resume_context)
    changes = _build_changes(job, material_view, resume_text)

    with connect(db_path) as conn:
        cursor = conn.execute(
            """
            INSERT INTO resume_versions (
                job_id, profile_id, evaluation_id, title, base_content
            ) VALUES (?, ?, ?, ?, ?)
            """,
            (
                job_id,
                profile["id"],
                evaluation["id"],
                title[:200],
                resume_text,
            ),
        )
        version_id = cursor.lastrowid
        for index, change in enumerate(changes, start=1):
            conn.execute(
                """
                INSERT INTO resume_changes (
                    version_id, change_type, section_key, before_text, after_text,
                    rationale, evidence_json, sort_order
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    version_id,
                    change["change_type"],
                    change["section_key"],
                    change["before_text"],
                    change["after_text"],
                    change["rationale"],
                    json_dump(change["evidence"]),
                    index,
                ),
            )
        _refresh_rendered_content(conn, version_id)

    version = get_resume_version(version_id, db_path)
    if version is None:
        raise RuntimeError("定制简历版本创建后无法读取")
    return version


def list_resume_versions(
    job_id: int,
    db_path: str | Path | None = None,
) -> list[dict[str, Any]]:
    with connect(db_path) as conn:
        job = conn.execute("SELECT id FROM jobs WHERE id = ?", (job_id,)).fetchone()
        if job is None:
            raise ValueError("岗位项目不存在")
        rows = conn.execute(
            """
            SELECT * FROM resume_versions
            WHERE job_id = ?
            ORDER BY id DESC
            """,
            (job_id,),
        ).fetchall()
        return [_version_response(row, conn, include_changes=False) for row in rows]


def get_resume_version(
    version_id: int,
    db_path: str | Path | None = None,
) -> dict[str, Any] | None:
    with connect(db_path) as conn:
        row = conn.execute(
            "SELECT * FROM resume_versions WHERE id = ?",
            (version_id,),
        ).fetchone()
        if row is None:
            return None
        return _version_response(row, conn, include_changes=True)


def update_resume_version(
    version_id: int,
    *,
    title: str | None = None,
    status: ResumeVersionStatus | None = None,
    template_id: ResumeTemplate | None = None,
    db_path: str | Path | None = None,
) -> dict[str, Any] | None:
    updates: list[str] = []
    values: list[Any] = []
    if title is not None:
        clean_title = title.strip()
        if not clean_title:
            raise ValueError("版本名称不能为空")
        updates.append("title = ?")
        values.append(clean_title[:200])
    if status is not None:
        if status == "final":
            current = get_resume_version(version_id, db_path)
            if current is None:
                return None
            gate = verify_candidate_material(
                str(current.get("rendered_content") or ""), db_path=db_path
            )
            if not gate["can_finalize"]:
                first = gate.get("issues", [{}])[0].get("message", "事实安全门未通过")
                raise ValueError(f"不能标记为可信定稿：{first}")
        updates.append("status = ?")
        values.append(status)
    if template_id is not None:
        if template_id not in _RESUME_TEMPLATES:
            raise ValueError("不支持的简历模板")
        updates.append("template_id = ?")
        values.append(template_id)
    with connect(db_path) as conn:
        existing = conn.execute(
            "SELECT id FROM resume_versions WHERE id = ?",
            (version_id,),
        ).fetchone()
        if existing is None:
            return None
        if updates:
            conn.execute(
                f"""
                UPDATE resume_versions
                SET {", ".join(updates)}, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
                """,
                (*values, version_id),
            )
    return get_resume_version(version_id, db_path)


def update_resume_change(
    version_id: int,
    change_id: int,
    *,
    decision: ResumeDecision | None = None,
    after_text: str | None = None,
    db_path: str | Path | None = None,
) -> dict[str, Any]:
    with connect(db_path) as conn:
        change = conn.execute(
            """
            SELECT id, after_text FROM resume_changes
            WHERE id = ? AND version_id = ?
            """,
            (change_id, version_id),
        ).fetchone()
        if change is None:
            raise ValueError("简历修改项不存在")

        updates: list[str] = []
        values: list[Any] = []
        if decision is not None:
            updates.append("decision = ?")
            values.append(decision)
        if after_text is not None:
            updates.extend(("after_text = ?", "user_edited = 1"))
            values.append(after_text.strip()[:100_000])
        if updates:
            conn.execute(
                f"""
                UPDATE resume_changes
                SET {", ".join(updates)}, updated_at = CURRENT_TIMESTAMP
                WHERE id = ? AND version_id = ?
                """,
                (*values, change_id, version_id),
            )
            _refresh_rendered_content(conn, version_id)

    version = get_resume_version(version_id, db_path)
    if version is None:
        raise ValueError("定制简历版本不存在")
    return version


def delete_resume_version(
    version_id: int,
    db_path: str | Path | None = None,
) -> bool:
    with connect(db_path) as conn:
        cursor = conn.execute("DELETE FROM resume_versions WHERE id = ?", (version_id,))
    return cursor.rowcount > 0


def export_resume_version(
    version_id: int,
    export_format: ResumeExportFormat,
    db_path: str | Path | None = None,
) -> tuple[bytes, str, str]:
    version = get_resume_version(version_id, db_path)
    if version is None:
        raise ValueError("定制简历版本不存在")
    content = str(version.get("rendered_content") or "").strip()
    if not content:
        raise ValueError("当前版本没有可导出的简历内容")
    gate = verify_candidate_material(content, db_path=db_path)
    if not gate["can_finalize"]:
        raise ValueError("事实安全门未通过，当前版本只能预览，不能导出为可信定稿")
    if export_format == "docx":
        payload = _build_docx(version["title"], content, version.get("template_id", "classic"))
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif export_format == "pdf":
        payload = _build_pdf(version["title"], content, version.get("template_id", "classic"))
        media_type = "application/pdf"
    else:
        raise ValueError("仅支持导出 DOCX 或 PDF")
    filename = f"{_safe_filename(version['title'])}.{export_format}"
    return payload, filename, media_type


def _version_response(
    row,
    conn,
    *,
    include_changes: bool,
) -> dict[str, Any]:
    version = row_to_dict(row) or {}
    decision_rows = conn.execute(
        """
        SELECT decision, COUNT(*) AS count
        FROM resume_changes
        WHERE version_id = ?
        GROUP BY decision
        """,
        (version["id"],),
    ).fetchall()
    counts = {"pending": 0, "accepted": 0, "rejected": 0}
    counts.update({item["decision"]: item["count"] for item in decision_rows})
    version["change_counts"] = counts
    version["change_count"] = sum(counts.values())
    if include_changes:
        rows = conn.execute(
            """
            SELECT * FROM resume_changes
            WHERE version_id = ?
            ORDER BY sort_order, id
            """,
            (version["id"],),
        ).fetchall()
        version["changes"] = [row_to_dict(item) for item in rows]
    else:
        version.pop("base_content", None)
        version.pop("rendered_content", None)
    return version


def _evaluation_material_view(
    evaluation: dict[str, Any], context: dict[str, Any]
) -> dict[str, Any]:
    facts = {int(item["id"]): item for item in context.get("confirmed_facts", [])}
    requirements = []
    for item in evaluation.get("effective_requirements", []):
        fact_ids = [int(value) for value in item.get("effective_fact_ids") or []]
        evidence = []
        for fact_id in fact_ids:
            fact = facts.get(fact_id)
            if not fact:
                continue
            statement = str(fact.get("statement") or "")
            evidence.append({
                "excerpt": statement,
                "matched_skills": extract_skills(statement),
                "matched_terms": [],
            })
        requirements.append({
            "id": item["requirement_key"], "text": item["text"],
            "importance": item["importance"],
            "status": item.get("effective_match_status", item["match_status"]),
            "evidence": evidence,
        })
    return {"requirements": requirements, "feedback": {}}


def _build_changes(
    job: dict[str, Any],
    analysis: dict[str, Any],
    resume_text: str,
) -> list[dict[str, Any]]:
    evidence: list[dict[str, str]] = []
    matched_terms: list[str] = []
    matched_term_keys: set[str] = set()
    prioritized_excerpts: list[str] = []
    feedback = analysis.get("feedback") or {}

    for requirement in analysis.get("requirements", []):
        effective_status = (
            feedback.get(requirement["id"], {}).get("status")
            or requirement.get("status")
        )
        if effective_status == "no_evidence":
            continue
        for item in requirement.get("evidence", []):
            excerpt = str(item.get("excerpt") or "").strip()
            if not excerpt:
                continue
            if excerpt not in prioritized_excerpts:
                prioritized_excerpts.append(excerpt)
            evidence.append(
                {
                    "source": "resume",
                    "requirement_id": requirement["id"],
                    "requirement": requirement["text"],
                    "excerpt": excerpt,
                }
            )
            for term in [
                *item.get("matched_skills", []),
                *item.get("matched_terms", []),
            ]:
                clean_term = str(term).strip()
                term_key = clean_term.casefold()
                if clean_term and term_key not in matched_term_keys:
                    matched_terms.append(clean_term)
                    matched_term_keys.add(term_key)

    job_title = str(job.get("job_title") or "目标岗位").strip()
    company_name = str(job.get("company_name") or "").strip()
    target_text = f"# {job_title}"
    if company_name:
        target_text += f"\n目标公司：{company_name}"

    if matched_terms:
        summary_text = (
            "## 职业概述\n"
            f"简历中可验证的岗位相关能力包括：{'、'.join(matched_terms[:8])}。"
            "以下内容仅重组和突出原简历已有事实。"
        )
        skills_text = "## 核心能力\n" + "\n".join(
            f"- {term}" for term in matched_terms[:12]
        )
    else:
        summary_text = (
            "## 职业概述\n"
            "当前版本仅重排原简历内容，未新增未经证实的能力表述。"
        )
        skills_text = ""

    original_lines = [
        line.strip()
        for line in resume_text.replace("\r\n", "\n").replace("\r", "\n").splitlines()
        if line.strip()
    ]
    priority_lines: list[str] = []
    remaining_lines: list[str] = []
    for line in original_lines:
        if any(line == excerpt or line.startswith(excerpt) for excerpt in prioritized_excerpts):
            if line not in priority_lines:
                priority_lines.append(line)
        else:
            remaining_lines.append(line)
    body_sections: list[str] = []
    if priority_lines:
        body_sections.append(
            "## 岗位相关经历与项目\n"
            + "\n".join(f"- {line}" for line in priority_lines)
        )
    if remaining_lines:
        body_sections.append(
            "## 其他经历与信息\n"
            + "\n".join(f"- {line}" for line in remaining_lines)
        )
    reordered_body = "\n\n".join(body_sections) or resume_text
    base_evidence = evidence or [
        {
            "source": "resume",
            "requirement_id": "",
            "requirement": "当前脱敏简历",
            "excerpt": resume_text[:280],
        }
    ]

    changes = [
        {
            "change_type": "target",
            "section_key": "target",
            "before_text": "",
            "after_text": target_text,
            "rationale": "明确该版本对应的目标岗位和公司",
            "evidence": [
                {
                    "source": "job",
                    "requirement_id": "",
                    "requirement": "岗位项目",
                    "excerpt": " · ".join(
                        item for item in (company_name, job_title) if item
                    ),
                }
            ],
        },
        {
            "change_type": "summary",
            "section_key": "summary",
            "before_text": "",
            "after_text": summary_text,
            "rationale": "只使用简历中已经出现且被岗位要求命中的能力生成概述",
            "evidence": base_evidence[:8],
        },
    ]
    if skills_text:
        changes.append(
            {
                "change_type": "skills",
                "section_key": "skills",
                "before_text": "",
                "after_text": skills_text,
                "rationale": "优先展示当前岗位能够从简历直接验证的能力关键词",
                "evidence": base_evidence[:12],
            }
        )
    changes.append(
        {
            "change_type": "reorder",
            "section_key": "body",
            "before_text": resume_text,
            "after_text": reordered_body,
            "rationale": "把与岗位要求直接相关的原始经历前置，其余内容保持原文",
            "evidence": base_evidence,
        }
    )
    return changes


def _refresh_rendered_content(conn, version_id: int) -> None:
    rows = conn.execute(
        """
        SELECT before_text, after_text, decision
        FROM resume_changes
        WHERE version_id = ?
        ORDER BY sort_order, id
        """,
        (version_id,),
    ).fetchall()
    sections = []
    for row in rows:
        content = row["before_text"] if row["decision"] == "rejected" else row["after_text"]
        if content.strip():
            sections.append(content.strip())
    conn.execute(
        """
        UPDATE resume_versions
        SET rendered_content = ?, status = 'draft', updated_at = CURRENT_TIMESTAMP
        WHERE id = ?
        """,
        ("\n\n".join(sections), version_id),
    )


def _build_docx(title: str, content: str, template_id: str = "classic") -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    compact = template_id == "compact"
    minimal = template_id == "minimal"
    section.top_margin = Cm(1.45 if compact else 1.8)
    section.bottom_margin = Cm(1.45 if compact else 1.8)
    section.left_margin = Cm(1.7 if compact else 2.0)
    section.right_margin = Cm(1.7 if compact else 2.0)
    section.start_type = WD_SECTION.NEW_PAGE

    font_name = _document_font_name()
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    for style_name, size, color in (
        ("Title", 20 if compact else 22, "1C2D3B" if minimal else "17324D"),
        ("Heading 1", 13 if compact else 15, "1C2D3B" if minimal else "17324D"),
        ("Heading 2", 11 if compact else 12, "4A5861" if minimal else "2D5B7D"),
        ("List Bullet", 9.7 if compact else 10.5, "263746"),
    ):
        style = styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    for line in content.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith("# "):
            paragraph = document.add_paragraph(text[2:], style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text.startswith("## "):
            document.add_heading(text[3:], level=1)
        elif text.startswith("- "):
            paragraph = document.add_paragraph(text[2:], style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
        else:
            paragraph = document.add_paragraph(text)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("BossCopilot 定制简历 · ")
    footer_run.font.name = font_name
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    page_run = footer.add_run()
    page_run.font.name = font_name
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_separator = OxmlElement("w:fldChar")
    field_separator.set(qn("w:fldCharType"), "separate")
    field_value = OxmlElement("w:t")
    field_value.text = "1"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    page_run._r.extend(
        (field_begin, instruction, field_separator, field_value, field_end)
    )
    document.core_properties.title = title
    document.core_properties.subject = "岗位定制简历"

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _build_pdf(title: str, content: str, template_id: str = "classic") -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("PDF 导出依赖未安装，请安装 reportlab") from exc

    font_name = "STSong-Light"
    for font_path in _pdf_font_candidates():
        if not font_path.exists():
            continue
        try:
            embedded_name = "BossCopilotResume"
            pdfmetrics.registerFont(TTFont(embedded_name, str(font_path)))
            font_name = embedded_name
            break
        except Exception:
            continue
    if font_name == "STSong-Light":
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    compact = template_id == "compact"
    minimal = template_id == "minimal"
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=(14 if compact else 18) * mm,
        leftMargin=(14 if compact else 18) * mm,
        topMargin=(14 if compact else 17) * mm,
        bottomMargin=(14 if compact else 17) * mm,
        title=title,
        author="BossCopilot",
    )
    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "ResumeBody",
        parent=base["BodyText"],
        fontName=font_name,
        fontSize=9.1 if compact else 9.8,
        leading=13.5 if compact else 15,
        textColor=colors.HexColor("#263746"),
        spaceAfter=5,
    )
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=body,
        fontSize=18 if compact else 20,
        leading=22 if compact else 25,
        textColor=colors.HexColor("#1C2D3B" if minimal else "#17324D"),
        spaceAfter=12,
    )
    heading = ParagraphStyle(
        "ResumeHeading",
        parent=body,
        fontSize=12 if compact else 13,
        leading=16 if compact else 18,
        textColor=colors.HexColor("#4A5861" if minimal else "#2D5B7D"),
        spaceBefore=7 if compact else 9,
        spaceAfter=5,
    )
    bullet = ParagraphStyle(
        "ResumeBullet",
        parent=body,
        leftIndent=11,
        firstLineIndent=-7,
        bulletIndent=2,
    )
    story = []
    for line in content.splitlines():
        text = line.strip()
        if not text:
            story.append(Spacer(1, 3))
        elif text.startswith("# "):
            story.append(Paragraph(escape(text[2:]), title_style))
        elif text.startswith("## "):
            story.append(Paragraph(escape(text[3:]), heading))
        elif text.startswith("- "):
            story.append(Paragraph(f"• {escape(text[2:])}", bullet))
        else:
            story.append(Paragraph(escape(text), body))

    def draw_footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor("#7A8793"))
        canvas.drawCentredString(A4[0] / 2, 9 * mm, f"BossCopilot 定制简历 · {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)
    return output.getvalue()


def _safe_filename(value: str) -> str:
    clean = "".join(
        character
        for character in value.strip()
        if character not in '<>:"/\\|?*\0'
    )
    return (clean or "定制简历")[:100]


def _document_font_name() -> str:
    if Path("/System/Library/Fonts/Hiragino Sans GB.ttc").exists():
        return "Hiragino Sans GB"
    if Path("/Library/Fonts/Arial Unicode.ttf").exists():
        return "Arial Unicode MS"
    if Path("C:/Windows/Fonts/msyh.ttc").exists():
        return "Microsoft YaHei"
    return "Noto Sans CJK SC"


def _pdf_font_candidates() -> tuple[Path, ...]:
    return (
        Path("/Library/Fonts/Arial Unicode.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
        Path("C:/Windows/Fonts/msyh.ttf"),
        Path("C:/Windows/Fonts/simhei.ttf"),
    )
