from __future__ import annotations

from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from dataclasses import dataclass, field
import re
import unicodedata

from docx import Document
from pypdf import PdfReader


MAX_RESUME_BYTES = 8 * 1024 * 1024
SUPPORTED_RESUME_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class ResumeParseResult:
    text: str
    parser: str
    warnings: list[str] = field(default_factory=list)


def parse_resume(filename: str, content: bytes) -> str:
    return parse_resume_result(filename, content).text


def parse_resume_result(filename: str, content: bytes, mode: str = "fast") -> ResumeParseResult:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_RESUME_SUFFIXES:
        raise ValueError("仅支持 PDF、Word（.docx）和文本简历")
    if not content:
        raise ValueError("简历文件为空")
    if len(content) > MAX_RESUME_BYTES:
        raise ValueError("简历文件不能超过 8MB")

    if mode not in {"fast", "enhanced"}:
        raise ValueError("解析模式必须是 fast 或 enhanced")

    warnings: list[str] = []
    if mode == "enhanced" and suffix in {".pdf", ".docx"}:
        try:
            text = _parse_with_docling(filename, content)
            parser = "docling"
        except Exception:
            text = _parse_lightweight(suffix, content)
            parser = "lightweight"
            warnings.append("增强解析暂不可用，已自动使用快速解析")
    else:
        text = _parse_lightweight(suffix, content)
        parser = "lightweight"

    normalized = normalize_resume_text(text)
    if len(normalized) < 20:
        raise ValueError("未能从简历中提取足够文字；扫描版 PDF 请使用增强解析")
    return ResumeParseResult(normalized[:100_000], parser, warnings)


def normalize_resume_text(text: str) -> str:
    """Keep resume text readable after PDF/Word font extraction.

    Some templates encode bullet points in the Unicode private-use area (for
    example ``\uf0b7``). Those glyphs are not meaningful text and otherwise end
    up in the editable resume. Convert line-leading private-use glyphs into a
    normal dash and discard other invisible/control characters without touching
    URLs, Chinese text, punctuation, or tables.
    """
    normalized = unicodedata.normalize("NFC", text).replace("\r\n", "\n").replace("\r", "\n")
    output: list[str] = []
    line_start = True
    for character in normalized:
        category = unicodedata.category(character)
        if character == "\n":
            output.append(character)
            line_start = True
            continue
        if character == "\t":
            output.append(character)
            line_start = False
            continue
        if category == "Co":
            if line_start:
                output.append("-")
            continue
        if category == "Cf" or category == "Cc":
            continue
        if character in {"•", "●", "▪", "◦", "■", "◆"} and line_start:
            output.append("-")
            continue
        output.append(character)
        line_start = character.isspace()

    return "\n".join(
        re.sub(r"^-\s*", "- ", line.strip())
        for line in "".join(output).splitlines()
        if line.strip()
    )


def _parse_lightweight(suffix: str, content: bytes) -> str:
    if suffix == ".pdf":
        return "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)
    if suffix == ".docx":
        document = Document(BytesIO(content))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            blocks.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(blocks)
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError:
        return content.decode("gb18030")


def _parse_with_docling(filename: str, content: bytes) -> str:
    # Docling is intentionally imported only for an explicit enhanced parse. It
    # is large and may need to initialize local OCR models on first use.
    from docling.document_converter import DocumentConverter

    with TemporaryDirectory(prefix="bosscopilot-resume-") as directory:
        path = Path(directory) / Path(filename).name
        path.write_bytes(content)
        result = DocumentConverter().convert(path)
        return result.document.export_to_markdown()
